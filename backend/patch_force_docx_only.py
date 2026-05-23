from pathlib import Path

p = Path("hwp_template_report_api.py")
text = p.read_text(encoding="utf-8")

helper = r'''

# ============================================================
# FORCE DOCX ONLY MODE
# - 입력: .hwp / .hwpx / .docx
# - 출력: 무조건 .docx
# - HWP/HWPX output 변환은 사용하지 않는다.
# ============================================================

FORCE_DOCX_ONLY = True

def force_extract_text_from_hwp(hwp_path: str | Path) -> str:
    hwp_path = Path(hwp_path).resolve()

    candidates = [
        ["hwp5txt", str(hwp_path)],
        ["python", "-m", "hwp5.hwp5txt", str(hwp_path)],
    ]

    last_err = ""

    for cmd in candidates:
        try:
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=120,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
            last_err = (result.stderr or result.stdout or "")[-1500:]
        except Exception as e:
            last_err = str(e)

    raise RuntimeError(
        "HWP 텍스트 추출 실패. 이 HWP는 서버에서 직접 읽지 못했습니다. "
        "가능하면 HWPX 또는 DOCX 양식으로 저장해서 업로드하세요. "
        f"detail={last_err}"
    )


def force_extract_text_from_docx(docx_path: str | Path) -> str:
    from docx import Document

    doc = Document(str(docx_path))
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def force_replace_docx_paragraph(paragraph, field_values: dict) -> int:
    original = paragraph.text or ""
    updated = original

    for field, value in field_values.items():
        value = str(value or "").strip()
        if not value:
            continue

        placeholders = [
            "{{" + field + "}}",
            "{{ " + field + " }}",
            "[[" + field + "]]",
            "[[ " + field + " ]]",
            "${" + field + "}",
            "${ " + field + " }",
        ]

        for ph in placeholders:
            updated = updated.replace(ph, value)

    if updated == original:
        return 0

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)

    return 1


def force_fill_docx_template(template_path: str | Path, output_path: str | Path, field_values: dict) -> int:
    from docx import Document

    template_path = Path(template_path).resolve()
    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))
    changed = 0

    for p in doc.paragraphs:
        changed += force_replace_docx_paragraph(p, field_values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    changed += force_replace_docx_paragraph(p, field_values)

    doc.save(str(output_path))
    return changed


def force_create_meeting_minutes_docx(output_path: str | Path, fields: dict) -> Path:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("회   의   록")
    r.bold = True
    r.font.size = Pt(20)

    doc.add_paragraph("")

    dept = doc.add_paragraph()
    dept.add_run("부 서 명 : ").bold = True
    dept.add_run(str(fields.get("부서명", "미기재") or "미기재"))

    doc.add_paragraph("")

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    rows = [
        ("건     명", fields.get("건명", "회의 내용 기반 요약")),
        ("진 행 시 간", fields.get("진행 시간", fields.get("진행시간", "    시      분부터\n    시      분까지"))),
        ("장     소", fields.get("장소", "미기재")),
        ("소집 및 발안자", fields.get("소집 및 발안자", "미기재")),
        ("의결 사항(요약)", fields.get("의결 사항(요약)", fields.get("의결사항", "미기재"))),
    ]

    for i, (label, value) in enumerate(rows):
        left = table.cell(i, 0)
        right = table.cell(i, 1)
        left.text = label
        right.text = str(value or "미기재")

        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for para in left.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("위 결의사항에 모두 동의하며 연서 및 날인함.")
    doc.add_paragraph("")

    participants = fields.get("참석자", [])
    if isinstance(participants, str):
        participants = [
            x.strip()
            for x in participants.replace(",", "\n").splitlines()
            if x.strip()
        ]

    if not isinstance(participants, list):
        participants = []

    for i in range(max(10, len(participants))):
        name = participants[i] if i < len(participants) else ""
        doc.add_paragraph(f" 참 석 자 : {name}                (인)")

    doc.save(str(output_path))
    return output_path


def force_get_template_text(template_path: str | Path) -> str:
    template_path = Path(template_path).resolve()
    ext = template_path.suffix.lower()

    if ext == ".docx":
        return force_extract_text_from_docx(template_path)

    if ext == ".hwp":
        return force_extract_text_from_hwp(template_path)

    if ext == ".hwpx":
        with tempfile.TemporaryDirectory(prefix="force_hwpx_read_") as tmp:
            work_dir = Path(tmp)
            extract_dir = work_dir / "unzipped"
            unzip_hwpx(template_path, extract_dir)
            section_files = iter_section_xml_files(extract_dir)
            paragraphs, _trees = load_paragraph_refs(section_files)
            return extract_visible_template_text(paragraphs)

    raise ValueError(f"지원하지 않는 템플릿 형식입니다: {ext}")


def force_generate_docx_only_report(
    template_path: str | Path,
    session_id: str,
    room_name: Optional[str],
    output_dir: str | Path,
    use_llm_field_detection: bool = True,
    user_prompt: str = "",
) -> Dict[str, Any]:
    template_path = Path(template_path).resolve()
    output_dir = Path(output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    transcript = read_session_transcript(room_name, session_id)
    if not transcript:
        raise ValueError("해당 세션의 회의 STT가 비어 있습니다.")

    document_context = build_document_context_for_minutes(session_id, transcript)

    template_text = force_get_template_text(template_path)

    fields = extract_placeholder_fields(template_text)

    if not fields and use_llm_field_detection:
        try:
            fields = detect_fields_with_qwen(template_text)
        except Exception as e:
            print(f"[WARN] force docx field detection failed: {e}")

    default_fields = globals().get("DEFAULT_MEETING_MINUTES_FIELDS", [
        "부서명",
        "건명",
        "장소",
        "소집 및 발안자",
        "의결 사항(요약)",
    ])

    merged_fields = []
    for f in list(fields or []) + list(default_fields):
        if f not in merged_fields:
            merged_fields.append(f)

    fields = merged_fields

    field_values = generate_field_values(
        fields,
        document_context,
        template_text=template_text,
        user_prompt=user_prompt,
    )

    for field in default_fields:
        if field not in field_values:
            field_values[field] = build_field_fallback_value(field, document_context)
            field_values[field] = force_field_value_by_type(field, field_values[field], document_context)

    stem = template_path.stem
    output_docx = output_dir / f"{stem}_filled_{session_id[:8]}.docx"

    changed = 0

    if template_path.suffix.lower() == ".docx":
        changed = force_fill_docx_template(template_path, output_docx, field_values)

    if template_path.suffix.lower() != ".docx" or changed == 0:
        force_create_meeting_minutes_docx(output_docx, field_values)
        changed = len(field_values)

    return {
        "ok": True,
        "templatePath": str(template_path),
        "roomName": room_name or "",
        "sessionId": session_id,
        "fields": fields,
        "fieldValues": field_values,
        "changedCount": changed,
        "outputPath": str(output_docx),
        "outputDocxPath": str(output_docx),
        "outputHwpxPath": "",
        "model": HWP_REPORT_MODEL,
        "mode": "force_docx_only",
    }

'''

if "def force_generate_docx_only_report" not in text:
    marker = "\n# ============================================================\n# 전체 파이프라인"
    if marker not in text:
        raise SystemExit("Cannot find pipeline marker")
    text = text.replace(marker, helper + marker, 1)

# generate_report_from_hwp_template 안에서 output_format 검증을 docx only로 완화/고정
old = '''    output_format = (output_format or "hwp").lower().strip()
    if output_format not in {"hwp", "hwpx", "docx"}:
        raise ValueError("output_format은 'hwp', 'hwpx', 'docx'만 가능합니다.")'''

new = '''    # 현재 서비스는 안정성을 위해 무조건 DOCX만 출력한다.
    output_format = "docx"'''

if old in text:
    text = text.replace(old, new, 1)
else:
    old2 = '''    output_format = (output_format or "hwp").lower().strip()
    if output_format not in {"hwp", "hwpx"}:
        raise ValueError("output_format은 'hwp' 또는 'hwpx'만 가능합니다.")'''
    if old2 in text:
        text = text.replace(old2, new, 1)

# out_dir 생성 직후 force return 삽입
old = '''    out_dir = Path(output_dir).resolve() if output_dir else OUTPUT_ROOT / datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir.mkdir(parents=True, exist_ok=True)
'''

new = '''    out_dir = Path(output_dir).resolve() if output_dir else OUTPUT_ROOT / datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir.mkdir(parents=True, exist_ok=True)

    # FORCE DOCX ONLY: hwp/hwpx/docx 입력 모두 docx로만 생성한다.
    return force_generate_docx_only_report(
        template_path=template_path,
        session_id=session_id,
        room_name=room_name,
        output_dir=out_dir,
        use_llm_field_detection=use_llm_field_detection,
        user_prompt=user_prompt,
    )
'''

if old not in text:
    print("[WARN] out_dir block not found, maybe already patched")
else:
    text = text.replace(old, new, 1)

# FastAPI 업로드 확장자 docx 허용
text = text.replace('if ext not in {".hwp", ".hwpx"}:', 'if ext not in {".hwp", ".hwpx", ".docx"}:')
text = text.replace("template은 .hwp 또는 .hwpx만 가능합니다.", "template은 .hwp, .hwpx 또는 .docx만 가능합니다.")

# 라우터 outputFormat 기본값 docx로
text = text.replace('outputFormat: str = Form("hwp")', 'outputFormat: str = Form("docx")')
text = text.replace('outputFormat: str = Form("hwpx")', 'outputFormat: str = Form("docx")')

# CLI choices docx만 사실상 허용
text = text.replace('choices=["hwp", "hwpx", "docx"]', 'choices=["docx"]')
text = text.replace('choices=["hwp", "hwpx"]', 'choices=["docx"]')
text = text.replace('default="hwp"', 'default="docx"')
text = text.replace('default="hwpx"', 'default="docx"')

# FileResponse media type docx 고정
text = text.replace(
    'media_type="application/octet-stream"',
    'media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"'
)

p.write_text(text, encoding="utf-8")
print("[OK] force DOCX-only patch applied")
