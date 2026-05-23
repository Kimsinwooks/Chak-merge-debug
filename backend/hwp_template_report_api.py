# -*- coding: utf-8 -*-
"""
hwp_template_report_api.py

목적
- 사용자가 업로드한 회의요약본 양식(.hwp 또는 .hwpx)을 HWPX로 변환/파싱한다.
- backend/data/{room_name}/sessions/{session_id}/live_recordings.sqlite3 에 저장된 회의 STT를 읽는다.
- 로컬 Ollama GEMMA(gemma3:27b 등)에 양식 항목별 작성 내용을 요청한다.
- 원본 HWP/HWPX 양식의 레이아웃을 최대한 유지하면서 텍스트 내용만 채운 HWP/HWPX 파일을 생성한다.

필수 환경
- Windows + 한글(HWP) 설치 + pywin32
- pip install pywin32 psutil requests fastapi python-multipart
- Ollama 실행: ollama serve
- 모델 설치: ollama pull gemma3:27b

FastAPI 연결
1) 이 파일을 backend/hwp_template_report_api.py 로 저장
2) backend/main.py 에 추가
   from hwp_template_report_api import router as hwp_template_report_router
   app.include_router(hwp_template_report_router)

CLI 테스트
python hwp_template_report_api.py --template "C:\\path\\회의요약본.hwp" --room "테스트룸" --session "SESSION_ID" --format hwp
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
import uuid
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import requests

try:
    import psutil
except Exception:  # pragma: no cover
    psutil = None

try:
    import winreg
    import win32com.client
except Exception:  # pragma: no cover
    winreg = None
    win32com = None

try:
    from fastapi import APIRouter, File, Form, HTTPException, UploadFile
    from fastapi.responses import FileResponse
except Exception:  # pragma: no cover
    APIRouter = None
    File = Form = UploadFile = None
    FileResponse = None


# ============================================================
# 기본 경로/모델 설정
# ============================================================

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
OUTPUT_ROOT = DATA_DIR / "filled_hwp_reports"

# HWP 양식 작성용 모델
# 프로젝트 내부 SLM_Loader의 qwen alias를 우선 사용한다.
HWP_REPORT_MODEL = (
    os.getenv("HWP_REPORT_MODEL")
    or os.getenv("REPORT_CHUNK_MODEL")
    or os.getenv("REALTIME_SLM_MODEL")
    or "Qwen2.5:3B"
)

try:
    from SLM_Loader import generate_ollama_chat_response
except Exception as e:
    generate_ollama_chat_response = None
    SLM_LOADER_IMPORT_ERROR = e
    
HWP_EXE_PATH = os.getenv(
    "HWP_EXE_PATH",
    r"C:\Program Files (x86)\Hnc\Office 2024\HOffice130\Bin\Hwp.exe",
)

RESERVED_ROOM_NAMES = {"_users", "sessions", "global_library"}
ROOM_NAME_PATTERN = re.compile(r"^[0-9A-Za-z가-힣ㄱ-ㅎㅏ-ㅣ _().-]{1,60}$")

COMMON_FIELD_HINTS = [
    "회의 제목", "회의명", "제목", "회의 일시", "일시", "장소", "참석자", "작성자",
    "회의 목적", "목적", "안건", "논의 안건", "주요 안건", "회의 내용", "주요 내용",
    "논의 내용", "결정 사항", "의사결정", "결론", "향후 계획", "후속 조치",
    "액션 아이템", "Action Item", "To-Do", "TODO", "담당자", "기한", "리스크", "특이사항",
]

PLACEHOLDER_PATTERNS = [
    re.compile(r"\{\{\s*([^{}\n]{1,40})\s*\}\}"),
    re.compile(r"\[\[\s*([^\[\]\n]{1,40})\s*\]\]"),
    re.compile(r"\$\{\s*([^{}\n]{1,40})\s*\}"),
]


@dataclass
class ParagraphRef:
    section_path: Path
    elem: Any
    text_nodes: List[Any]
    text: str
    index: int


# ============================================================
# HWP COM 유틸
# ============================================================

def kill_hwp_processes() -> None:
    """실행 중인 HWP 프로세스를 종료한다. COM 변환 충돌 방지용."""
    if psutil is None:
        return
    for proc in psutil.process_iter(["name"]):
        try:
            name = proc.info.get("name") or ""
            if "hwp" in name.lower():
                print(f"[HWP] 실행 중인 HWP 종료: {name}")
                proc.terminate()
                proc.wait(timeout=5)
        except Exception as e:
            print(f"[WARN] HWP 프로세스 종료 실패: {e}")


def check_hwp_com_registration() -> bool:
    """HWPFrame.HwpObject COM 등록 여부를 확인하고, 없으면 /regserver 등록을 시도한다."""
    if winreg is None:
        raise RuntimeError("winreg를 사용할 수 없습니다. Windows 환경에서 실행하세요.")
    try:
        key = winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, r"HWPFrame.HwpObject\CLSID")
        winreg.CloseKey(key)
        return True
    except FileNotFoundError:
        if not Path(HWP_EXE_PATH).exists():
            raise RuntimeError(f"HWP 실행 파일을 찾을 수 없습니다: {HWP_EXE_PATH}")
        subprocess.run([HWP_EXE_PATH, "/regserver"], check=False)
        return False


def _new_hwp_object():
    if win32com is None:
        raise RuntimeError("pywin32가 필요합니다. `pip install pywin32` 후 Windows에서 실행하세요.")
    check_hwp_com_registration()
    hwp = win32com.client.DispatchEx("HWPFrame.HwpObject")
    try:
        hwp.RegisterModule("FilePathCheckDLL", "SecurityModule")
    except Exception:
        # 일부 환경에서는 보안 모듈 등록이 실패해도 Open/SaveAs가 동작한다.
        pass
    return hwp


def convert_hwp_to_hwpx(hwp_path: str | Path, hwpx_path: str | Path) -> Path:
    hwp_path = Path(hwp_path).resolve()
    hwpx_path = Path(hwpx_path).resolve()
    hwpx_path.parent.mkdir(parents=True, exist_ok=True)

    hwp = _new_hwp_object()
    try:
        hwp.Open(str(hwp_path), "HWP", "")
        hwp.SaveAs(str(hwpx_path), "HWPX", "")
    finally:
        try:
            hwp.Quit()
        except Exception:
            pass
    return hwpx_path


def convert_hwpx_to_hwp(hwpx_path: str | Path, hwp_path: str | Path) -> Path:
    hwpx_path = Path(hwpx_path).resolve()
    hwp_path = Path(hwp_path).resolve()
    hwp_path.parent.mkdir(parents=True, exist_ok=True)

    hwp = _new_hwp_object()
    try:
        hwp.Open(str(hwpx_path), "HWPX", "")
        hwp.SaveAs(str(hwp_path), "HWP", "")
    finally:
        try:
            hwp.Quit()
        except Exception:
            pass
    return hwp_path


# ============================================================
# 룸/세션 transcript 읽기
# ============================================================

def sanitize_room_name(room_name: str) -> str:
    room_name = (room_name or "").strip()
    if not room_name:
        raise ValueError("room_name이 필요합니다.")
    if room_name in RESERVED_ROOM_NAMES:
        raise ValueError(f"예약된 룸 이름은 사용할 수 없습니다: {room_name}")
    if "/" in room_name or "\\" in room_name or ".." in room_name:
        raise ValueError("room_name에 경로 문자를 사용할 수 없습니다.")
    if not ROOM_NAME_PATTERN.match(room_name):
        raise ValueError("room_name은 한글, 영어, 숫자, 공백, _, -, ., 괄호만 사용할 수 있습니다.")
    return room_name


def validate_session_id(session_id: str) -> str:
    session_id = (session_id or "").strip()
    if not session_id or "/" in session_id or "\\" in session_id or ".." in session_id:
        raise ValueError("잘못된 session_id입니다.")
    return session_id


def get_live_db_path(room_name: str, session_id: str) -> Path:
    room_name = sanitize_room_name(room_name)
    session_id = validate_session_id(session_id)
    return DATA_DIR / room_name / "sessions" / session_id / "live_recordings.sqlite3"


def find_live_db_path_by_session_id(session_id: str) -> Optional[Path]:
    """room_name을 모를 때 backend/data/*/sessions/{session_id}/live_recordings.sqlite3를 검색."""
    session_id = validate_session_id(session_id)
    if not DATA_DIR.exists():
        return None
    for room_dir in DATA_DIR.iterdir():
        if not room_dir.is_dir():
            continue
        if room_dir.name.startswith("_") or room_dir.name in RESERVED_ROOM_NAMES:
            continue
        candidate = room_dir / "sessions" / session_id / "live_recordings.sqlite3"
        if candidate.exists():
            return candidate
    return None


def _format_sec(sec: float | int | None) -> str:
    sec_int = int(sec or 0)
    mm = sec_int // 60
    ss = sec_int % 60
    return f"{mm:02d}:{ss:02d}"


def read_session_transcript(room_name: Optional[str], session_id: str) -> str:
    """
    현재 프로젝트의 실제 저장 구조를 우선 사용한다.

    1순위: backend/data/meeting_app.sqlite3 의 transcript_lines
    2순위: backend/data/{room}/sessions/{session_id}/live_recordings.sqlite3 의 live_transcripts
    """
    session_id = validate_session_id(session_id)

    # ============================================================
    # 1. 현재 실제 저장 구조: meeting_app.sqlite3 / transcript_lines
    # ============================================================
    meeting_app_db = DATA_DIR / "meeting_app.sqlite3"

    if meeting_app_db.exists():
        conn = sqlite3.connect(str(meeting_app_db))
        conn.row_factory = sqlite3.Row

        try:
            rows = conn.execute(
                """
                SELECT speaker, start_sec, end_sec, text, created_at
                FROM transcript_lines
                WHERE session_id = ?
                ORDER BY start_sec ASC, end_sec ASC, created_at ASC
                """,
                (session_id,),
            ).fetchall()
        finally:
            conn.close()

        lines = []
        for row in rows:
            text = (row["text"] or "").strip()
            if not text:
                continue

            speaker = row["speaker"] or "익명1"
            start = _format_sec(row["start_sec"])
            end = _format_sec(row["end_sec"] or row["start_sec"])
            lines.append(f"[{start}~{end}] {speaker}: {text}")

        if lines:
            return "\n".join(lines).strip()

    # ============================================================
    # 2. 예전 구조 fallback: live_recordings.sqlite3 / live_transcripts
    # ============================================================
    if room_name:
        db_path = get_live_db_path(room_name, session_id)
    else:
        db_path = find_live_db_path_by_session_id(session_id)

    if db_path is None or not db_path.exists():
        raise FileNotFoundError(
            "회의 STT를 찾지 못했습니다. "
            f"meeting_app.sqlite3 transcript_lines와 live_recordings.sqlite3 모두 비어 있습니다. "
            f"session_id={session_id}, room_name={room_name or ''}"
        )

    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row

    try:
        rows = conn.execute(
            """
            SELECT speaker, start_sec, end_sec, text, created_at
            FROM live_transcripts
            WHERE session_id = ?
            ORDER BY start_sec ASC, created_at ASC
            """,
            (session_id,),
        ).fetchall()
    finally:
        conn.close()

    lines = []
    for row in rows:
        text = (row["text"] or "").strip()
        if not text:
            continue

        speaker = row["speaker"] or "익명1"
        start = _format_sec(row["start_sec"])
        end = _format_sec(row["end_sec"] or row["start_sec"])
        lines.append(f"[{start}~{end}] {speaker}: {text}")

    return "\n".join(lines).strip()


# ============================================================
# Ollama GEMMA 호출
# ============================================================

def extract_json_object(text: str) -> dict:
    cleaned = (text or "").strip()
    cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start < 0 or end <= start:
        raise ValueError(f"JSON object를 찾을 수 없습니다: {cleaned[:500]}")
    return json.loads(cleaned[start : end + 1])


def call_qwen_chat_json(
    system_prompt: str,
    user_prompt: str,
    model_name: str = HWP_REPORT_MODEL,
    timeout: int = 900,
    temperature: float = 0.1,
    num_predict: int = 4096,
) -> dict:
    """
    프로젝트 내부 SLM_Loader를 통해 Qwen/Ollama를 호출하고 JSON 객체로 파싱한다.
    직접 requests.post('/api/chat')를 호출하지 않는다.
    """
    if generate_ollama_chat_response is None:
        raise RuntimeError(
            f"SLM_Loader.generate_ollama_chat_response import 실패: "
            f"{SML_LOADER_IMPORT_ERROR if 'SML_LOADER_IMPORT_ERROR' in globals() else SLM_LOADER_IMPORT_ERROR}"
        )

    content = generate_ollama_chat_response(
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        model_name=model_name,
        max_new_tokens=num_predict,
        temperature=temperature,
        top_p=0.9,
        format_json=True,
    )

    if not content or not str(content).strip():
        raise RuntimeError("Qwen 응답이 비어 있습니다.")

    return extract_json_object(str(content))

def call_qwen_chat_text(
    system_prompt: str,
    user_prompt: str,
    model_name: str = HWP_REPORT_MODEL,
    temperature: float = 0.0,
    num_predict: int = 1024,
) -> str:
    """
    Qwen에게 JSON이 아니라 일반 텍스트 답변을 요청한다.
    HWP 필드값 생성은 JSON key 매칭보다 필드별 텍스트 생성이 안정적이다.
    """
    if generate_ollama_chat_response is None:
        raise RuntimeError(
            f"SLM_Loader.generate_ollama_chat_response import 실패: "
            f"{SML_LOADER_IMPORT_ERROR if 'SML_LOADER_IMPORT_ERROR' in globals() else SLM_LOADER_IMPORT_ERROR}"
        )

    content = generate_ollama_chat_response(
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        model_name=model_name,
        max_new_tokens=num_predict,
        temperature=temperature,
        top_p=0.9,
        format_json=False,
    )

    return str(content or "").strip()

def detect_fields_with_qwen(template_text: str) -> List[str]:
    system = """
너는 HWP 회의요약본 양식에서 사용자가 채워야 하는 항목명을 추출하는 파서다.
반드시 JSON 하나만 출력한다. 설명하지 마라.
""".strip()

    user = f"""
아래는 회의요약본 양식에서 추출한 텍스트다.

사용자가 회의 내용을 기반으로 작성해야 하는 항목명만 추출하라.
로고, 회사명, 고정 안내문, 문서번호, 페이지 번호처럼 채우는 항목이 아닌 것은 제외하라.

출력 JSON schema:
{{"fields": ["회의 제목", "회의 내용", "향후 계획"]}}

[양식 텍스트]
{template_text[:12000]}
""".strip()

    data = call_qwen_chat_json(
        system,
        user,
        model_name=HWP_REPORT_MODEL,
        num_predict=1024,
        temperature=0.1,
    )

    fields = data.get("fields", [])
    if not isinstance(fields, list):
        return []

    return normalize_fields([str(x).strip() for x in fields if str(x).strip()])

def normalize_field_key(value: str) -> str:
    """
    필드명 비교용 정규화.
    예:
    '의결 사항(요약)' -> '의결사항요약'
    '소집 및 발안자' -> '소집및발안자'
    """
    return re.sub(r"[^0-9a-zA-Z가-힣]", "", str(value or "")).lower()


def get_value_by_fuzzy_key(data: dict, field: str) -> str:
    """
    Qwen이 JSON key를 살짝 다르게 내보내도 최대한 원래 필드와 매칭한다.
    """
    if field in data:
        return data.get(field, "")

    target = normalize_field_key(field)

    for key, value in data.items():
        if normalize_field_key(key) == target:
            return value

    for key, value in data.items():
        nk = normalize_field_key(key)
        if target in nk or nk in target:
            return value

    return ""

def clean_transcript_for_document(transcript: str) -> str:
    """
    HWP fallback에 STT 원문이 그대로 들어가지 않도록
    시간표시, 화자표시, 너무 짧은 발화를 제거한다.
    """
    cleaned_lines = []

    for line in str(transcript or "").splitlines():
        line = line.strip()
        if not line:
            continue

        # [00:01~00:03] 제거
        line = re.sub(r"\[\d{1,2}:\d{2}(?::\d{2})?\s*~\s*\d{1,2}:\d{2}(?::\d{2})?\]", "", line)

        # 익명1: 제거
        line = re.sub(r"^\s*익명\d+\s*:\s*", "", line)

        # 너무 짧거나 의미 없는 발화 제거
        if len(line) < 8:
            continue
        if line in {"어...", "음...", "네", "아", "어"}:
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


def simple_fallback_summary(transcript: str, max_items: int = 5) -> str:
    """
    Qwen이 실패했을 때 의결 사항 칸에 넣을 최소 요약.
    원문 전체를 넣지 않고, 의미 있는 문장 몇 개만 공식 문서체처럼 정리한다.
    """
    cleaned = clean_transcript_for_document(transcript)
    if not cleaned.strip():
        return "회의 내용이 명확히 기록되지 않아 세부 의결 사항은 미기재함."

    sentences = re.split(r"(?<=[.!?。])\s+|\n+", cleaned)
    sentences = [s.strip() for s in sentences if len(s.strip()) >= 10]

    selected = sentences[:max_items]

    if not selected:
        return "회의 주요 내용에 대한 검토를 진행하였으며, 세부 사항은 후속 회의에서 추가 확인하기로 함."

    formalized = []
    for s in selected:
        s = str(s).strip()
        s = re.sub(r"^(아니|근데|그래서|음|어|네)\s*", "", s)
        s = s.replace("요", "").strip()
        s = s.rstrip(".。")
        if not s:
            continue
        if not s.endswith(("함", "됨", "예정", "필요")):
            s = s + "에 대해 논의함"
        formalized.append(f"- {s[:180]}")

    if not formalized:
        return "회의 주요 내용에 대한 검토를 진행하였으며, 세부 사항은 후속 회의에서 추가 확인하기로 함."

    return "\n".join(formalized)


def build_fallback_field_values(fields: List[str], transcript: str) -> Dict[str, str]:
    """
    Qwen이 필드명과 정확히 맞는 JSON을 만들지 못했을 때도
    HWP 생성이 실패하지 않도록 최소값을 채우는 fallback.
    """
    summary = simple_fallback_summary(transcript)

    values: Dict[str, str] = {}

    for field in fields:
        key = normalize_field_key(field)

        if "부서" in key:
            values[field] = "미기재"

        elif "장소" in key:
            values[field] = "미기재"

        elif "소집" in key or "발안자" in key or "작성자" in key:
            values[field] = "미기재"

        elif "건명" in key or "제목" in key or "회의명" in key:
            values[field] = "회의 내용 기반 요약"

        elif "의결" in key or "요약" in key or "내용" in key or "결정" in key:
            values[field] = summary

        elif "참석" in key:
            values[field] = "미기재"

        else:
            values[field] = summary

    return values

def clean_qwen_field_answer(text: str) -> str:
    """
    Qwen 응답에서 코드블록, 따옴표, 이상한 JSON 대화 구조를 제거한다.
    """
    raw = str(text or "").strip()

    if not raw:
        return ""

    raw = re.sub(r"^```(?:json)?", "", raw).strip()
    raw = re.sub(r"```$", "", raw).strip()

    # Qwen이 또 user/system/assistant JSON을 만들면 해당 응답은 버린다.
    try:
        obj = extract_json_object(raw)
        if isinstance(obj, dict):
            bad_keys = {"user", "system", "assistant", "responses", "questions", "instructions"}
            if any(k in obj for k in bad_keys):
                return ""

            # 혹시 {"answer": "..."} 형태면 answer만 사용
            for key in ["answer", "value", "content", "text", "result"]:
                if key in obj and isinstance(obj[key], str):
                    return obj[key].strip()

            # 문자열 값이 하나만 있으면 그 값 사용
            string_values = [v.strip() for v in obj.values() if isinstance(v, str) and v.strip()]
            if len(string_values) == 1:
                return string_values[0]

    except Exception:
        pass

    raw = raw.strip().strip('"').strip("'").strip()

    # 불필요한 접두어 제거
    raw = re.sub(r"^(답변|작성 내용|결과|요약)\s*[:：]\s*", "", raw).strip()

    return raw



def extract_minutes_bullets_from_context(context: str, max_items: int = 5) -> str:
    """
    cached report context에서 회의록에 넣기 좋은 요약/주제 문장을 우선 추출한다.
    원본 STT 말투를 그대로 쓰지 않고 공식 회의록 문체로 정리한다.
    """
    text = str(context or "").strip()
    if not text:
        return "회의 주요 내용에 대한 검토를 진행하였으며, 세부 사항은 후속 회의에서 추가 확인하기로 함."

    candidates = []

    # report flatten 결과에서 summary/description/title 이후 내용을 우선 수집
    lines = [x.strip() for x in text.splitlines() if x.strip()]
    for i, line in enumerate(lines):
        low = line.lower().strip("# ").strip()
        if low in {"summary", "overallsummary", "overall_summary", "description", "title"}:
            if i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if len(nxt) >= 15:
                    candidates.append(nxt)

    # 그래도 부족하면 문장 단위 추출
    if len(candidates) < 3:
        import re
        extra = re.split(r"(?<=[.!?。])\s+|\n+", text)
        for s in extra:
            s = s.strip()
            if len(s) >= 25 and not s.startswith("##"):
                candidates.append(s)

    cleaned = []
    banned = [
        "또 다른 집에서", "아마라이", "스페인", "런던",
        "스터디언니아", "흥미로웠던 곳", "하와이"
    ]

    for s in candidates:
        if any(b in s for b in banned):
            continue

        s = s.strip()
        s = s.replace("습니다.", "음.")
        s = s.rstrip(".。")
        s = s[:180]

        if not s:
            continue

        if not s.endswith(("함", "됨", "예정", "필요", "확인")):
            s = s + "에 대해 논의함"

        bullet = f"- {s}"
        if bullet not in cleaned:
            cleaned.append(bullet)

        if len(cleaned) >= max_items:
            break

    if not cleaned:
        return "회의 주요 내용에 대한 검토를 진행하였으며, 세부 사항은 후속 회의에서 추가 확인하기로 함."

    return "\n".join(cleaned)


def build_field_fallback_value(field: str, transcript: str) -> str:
    """
    Qwen이 특정 필드 생성을 실패했을 때 필드별 최소값을 만든다.
    """
    key = normalize_field_key(field)

    cleaned = clean_transcript_for_document(transcript)
    summary = simple_fallback_summary(transcript)

    if "부서" in key:
        return "미기재"

    if "장소" in key:
        return "미기재"

    if "소집" in key or "발안자" in key or "작성자" in key:
        return "미기재"

    if "건명" in key or "제목" in key or "회의명" in key:
        # 너무 긴 STT가 제목으로 들어가지 않도록 고정된 최소 제목 사용
        if "충전" in cleaned:
            return "충전 관련 회의"
        if "시험" in cleaned or "세미나" in cleaned:
            return "시험 준비 및 세미나 일정 관련 회의"
        return "회의 내용 기반 요약"

    if "참석" in key:
        return "미기재"

    if "의결" in key or "요약" in key or "내용" in key or "결정" in key:
        return extract_minutes_bullets_from_context(transcript)

    return summary

def force_field_value_by_type(field: str, value: str, transcript: str) -> str:
    """
    Qwen이 필드 성격에 맞지 않는 값을 생성했을 때 강제로 보정한다.
    부서명/장소/소집자 같은 메타데이터 칸에 회의 요약이 들어가는 문제를 막는다.
    """
    key = normalize_field_key(field)
    value = str(value or "").strip()

    # Qwen이 만든 불필요한 마크다운/항목기호 제거
    value = re.sub(r"^[-•]\s*", "", value).strip()
    value = value.replace("\r\n", "\n").replace("\r", "\n")

    # 1. 부서명: STT에 명확한 부서명이 없으면 미기재
    if "부서" in key:
        # 회의 STT에서 부서명이 명시되는 경우만 인정
        m = re.search(r"(부서명|부서)\s*[:：]\s*([가-힣A-Za-z0-9\s]{2,30})", transcript or "")
        if m:
            return m.group(2).strip()[:30]
        return "미기재"

    # 2. 장소: STT에 명확한 장소가 없으면 미기재
    if "장소" in key:
        m = re.search(r"(장소|회의장소)\s*[:：]\s*([가-힣A-Za-z0-9\s]{2,40})", transcript or "")
        if m:
            return m.group(2).strip()[:40]
        return "미기재"

    # 3. 소집 및 발안자: STT에 명확한 이름이 없으면 미기재
    if "소집" in key or "발안자" in key or "작성자" in key:
        m = re.search(r"(소집자|발안자|작성자)\s*[:：]\s*([가-힣A-Za-z0-9\s]{2,30})", transcript or "")
        if m:
            return m.group(2).strip()[:30]
        return "미기재"

    # 4. 건명/제목: 무조건 한 줄만 허용
    if "건명" in key or "제목" in key or "회의명" in key:
        one_line = re.sub(r"\s+", " ", value).strip()
        one_line = re.sub(r"^[-•]\s*", "", one_line).strip()

        # Qwen이 요약문을 길게 넣으면 fallback 제목으로 대체
        if not one_line or len(one_line) > 60 or "\n" in value:
            cleaned = clean_transcript_for_document(transcript)
            if "시험" in cleaned and "세미나" in cleaned:
                return "시험 준비 및 세미나 일정 관련 회의"
            if "시각화" in cleaned:
                return "시각화 기능 개선 관련 회의"
            if "충전" in cleaned:
                return "충전 상태 확인 관련 회의"
            return "회의 내용 기반 요약"

        return one_line[:60]

    # 5. 의결/요약/내용: 여러 줄 허용
    if "의결" in key or "요약" in key or "내용" in key or "결정" in key:
        if not value:
            return simple_fallback_summary(transcript)

        # 너무 긴 경우 잘라냄
        lines = [line.strip() for line in value.splitlines() if line.strip()]
        lines = lines[:6]

        if not lines:
            return simple_fallback_summary(transcript)

        return "\n".join(lines)[:1200]

    # 6. 기타 필드
    if not value:
        return "미기재"

    return value[:500]

def generate_field_values(
    fields: List[str],
    transcript: str,
    template_text: str = "",
    user_prompt: str = "",
) -> Dict[str, str]:
    """
    기존 방식:
    - Qwen에게 전체 필드 JSON을 한 번에 요청
    - Qwen이 key를 못 맞추면 실패

    수정 방식:
    - 필드 하나씩 Qwen에게 plain text로 요청
    - JSON key 매칭 문제 제거
    - 실패한 필드만 fallback 사용
    """
    if not fields:
            raise ValueError("채울 필드가 없습니다.")
    if not transcript:
        raise ValueError("회의 transcript가 비어 있습니다.")

    user_prompt = str(user_prompt or "").strip()[:2000]
    extra_instruction = user_prompt if user_prompt else "없음"

    values: Dict[str, str] = {}

    for field in fields:
        key = normalize_field_key(field)

        system = """
너는 회사 공식 회의요약본을 작성하는 로컬 회의록 작성 AI다.

규칙:
- 사용자가 요청한 항목 하나에 들어갈 내용만 작성한다.
- JSON을 출력하지 않는다.
- 마크다운을 출력하지 않는다.
- 코드블록을 출력하지 않는다.
- 항목명은 반복하지 않는다.
- 회의 STT에 없는 사실은 만들지 않는다.
- 알 수 없는 값은 "미기재"라고 쓴다.
- 공식 문서체 한국어로 작성한다.
""".strip()

        if "부서" in key:
            instruction = '부서명을 작성하라. STT에 명확한 부서명이 없으면 "미기재"라고만 작성하라.'
            max_tokens = 80

        elif "건명" in key or "제목" in key or "회의명" in key:
            instruction = "회의의 핵심 주제를 20자 내외의 한 줄 제목으로 작성하라."
            max_tokens = 120

        elif "장소" in key:
            instruction = '회의 장소를 작성하라. STT에 명확한 장소가 없으면 "미기재"라고만 작성하라.'
            max_tokens = 80

        elif "소집" in key or "발안자" in key or "작성자" in key:
            instruction = '소집 및 발안자를 작성하라. STT에 명확한 인물이 없으면 "미기재"라고만 작성하라.'
            max_tokens = 100

        elif "참석" in key:
            instruction = '참석자를 작성하라. STT에 명확한 참석자가 없으면 "미기재"라고만 작성하라.'
            max_tokens = 120

        elif "의결" in key or "요약" in key or "내용" in key or "결정" in key:
            instruction = """
회의의 의결 사항 또는 요약을 작성하라.
반드시 공식 회의록 문체로 작성하라.
구어체, 반말, 감탄사, STT 원문 그대로의 문장을 절대 복사하지 마라.
결정 사항, 논의 결과, 후속 조치 중심으로 3~5개 항목으로 정리하라.
각 항목은 "- "로 시작하라.
예:
- 데이터 전처리 방식에 대한 검토를 완료하고, 수정이 필요한 항목을 재점검하기로 함.
- 모델 학습 결과 공유 및 추가 실험 계획을 다음 회의에서 확인하기로 함.
""".strip()
            max_tokens = 700

        else:
            instruction = f'"{field}" 항목에 들어갈 내용을 회의 STT 기반으로 작성하라.'
            max_tokens = 400

        user = f"""
작성할 항목명:
{field}

작성 지시:
{instruction}

사용자 추가 작성 지시:
{extra_instruction}

사용자 추가 작성 지시 반영 규칙:
- 위 지시가 "없음"이 아니면 문체, 강조점, 작성 방향에 반영한다.
- 단, 회의 STT에 없는 사실을 새로 만들면 안 된다.
- 사용자 지시가 "없음"이면 기존 방식대로 작성한다.

양식 전체 텍스트 참고:
{template_text[:2500] if template_text else "(없음)"}

회의 STT:
{transcript[-16000:]}

위 항목에 들어갈 내용만 출력하라.
""".strip()

        try:
            raw = call_qwen_chat_text(
                system,
                user,
                model_name=HWP_REPORT_MODEL,
                num_predict=max_tokens,
                temperature=0.0,
            )
            value = clean_qwen_field_answer(raw)
        except Exception as e:
            print(f"[WARN] Qwen 필드 생성 실패 field={field}: {e}")
            value = ""

        if not value:
            value = build_field_fallback_value(field, transcript)

        # 너무 긴 값 방지
        if "의결" in key or "요약" in key or "내용" in key or "결정" in key:
            value = value[:1200]
        else:
            value = value[:180]

        value = force_field_value_by_type(field, value, transcript)
        values[field] = value.strip() or "미기재"

    return values



# ============================================================
# Report context reader
# - 문서 생성 품질 향상을 위해 원본 STT보다
#   meeting_report_cache / report 계열 테이블의 정제된 회의록을 우선 사용한다.
# ============================================================

def read_cached_meeting_report_text(session_id: str) -> str:
    """
    회의록 정리 탭에 표시되는 정제된 report/cache 내용을 최대한 읽는다.
    테이블/컬럼명이 버전별로 다를 수 있으므로 유연하게 탐색한다.
    """
    session_id = validate_session_id(session_id)
    meeting_app_db = DATA_DIR / "meeting_app.sqlite3"

    if not meeting_app_db.exists():
        return ""

    candidate_tables = [
        "meeting_report_cache",
        "meeting_reports",
        "meeting_analysis",
        "meeting_sessions",
    ]

    text_columns_priority = [
        "report_markdown",
        "report_md",
        "markdown",
        "minutes",
        "meeting_minutes",
        "report",
        "report_text",
        "analysis_text",
        "summary",
        "summary_text",
        "content",
        "result",
        "raw_text",
    ]

    json_columns_priority = [
        "report_json",
        "analysis_json",
        "result_json",
        "cache_json",
        "raw_json",
        "data",
    ]

    try:
        conn = sqlite3.connect(str(meeting_app_db))
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()

        existing_tables = {
            row[0]
            for row in cur.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
        }

        for table in candidate_tables:
            if table not in existing_tables:
                continue

            cols = [r[1] for r in cur.execute(f"PRAGMA table_info({table})").fetchall()]
            if "session_id" not in cols:
                continue

            row = cur.execute(
                f"SELECT * FROM {table} WHERE session_id = ? ORDER BY rowid DESC LIMIT 1",
                (session_id,),
            ).fetchone()

            if not row:
                continue

            # 1) 일반 텍스트/마크다운 컬럼 우선
            for col in text_columns_priority:
                if col in cols:
                    value = row[col]
                    if value and str(value).strip():
                        return str(value).strip()

            # 2) JSON 컬럼에서 회의록/요약 필드 추출
            for col in json_columns_priority:
                if col not in cols:
                    continue

                raw = row[col]
                if not raw or not str(raw).strip():
                    continue

                try:
                    obj = json.loads(str(raw))
                except Exception:
                    continue

                extracted = flatten_report_json_to_text(obj)
                if extracted.strip():
                    return extracted.strip()

    except Exception as e:
        print(f"[WARN] read_cached_meeting_report_text failed: {e}")
        return ""

    finally:
        try:
            conn.close()
        except Exception:
            pass

    return ""


def flatten_report_json_to_text(obj: Any) -> str:
    """
    report_json 구조가 일정하지 않아도 회의록/요약/결정사항 중심으로 텍스트화한다.
    """
    if obj is None:
        return ""

    if isinstance(obj, str):
        return obj.strip()

    parts = []

    if isinstance(obj, dict):
        preferred_keys = [
            "minutes",
            "meetingMinutes",
            "meeting_minutes",
            "report",
            "reportMarkdown",
            "report_markdown",
            "summary",
            "overallSummary",
            "overall_summary",
            "decisions",
            "decisionItems",
            "todos",
            "actionItems",
            "topicBlocks",
            "topics",
        ]

        for key in preferred_keys:
            if key in obj:
                value = flatten_report_json_to_text(obj.get(key))
                if value:
                    parts.append(f"## {key}\n{value}")

        # preferred key에서 아무것도 못 뽑으면 전체 dict 순회
        if not parts:
            for key, value in obj.items():
                value_text = flatten_report_json_to_text(value)
                if value_text:
                    parts.append(f"## {key}\n{value_text}")

    elif isinstance(obj, list):
        for item in obj:
            value_text = flatten_report_json_to_text(item)
            if value_text:
                parts.append(value_text)

    else:
        return str(obj).strip()

    return "\n\n".join(parts).strip()


def build_document_context_for_minutes(session_id: str, transcript: str) -> str:
    """
    문서 생성용 context.
    1순위: 회의록 정리 탭에 해당하는 cached report
    2순위: 원본 transcript
    """
    report_text = read_cached_meeting_report_text(session_id)

    if report_text and len(report_text.strip()) >= 30:
        print(f"[HWP_DOC] using cached meeting report context, chars={len(report_text)}")
        return report_text.strip()

    print("[HWP_DOC] cached meeting report not found. fallback to transcript.")
    return transcript


# ============================================================
# HWP input -> DOCX fallback
# - Ubuntu/Linux에서 .hwp 원본 서식을 직접 치환하기는 어렵다.
# - 대신 hwp5txt/pyhwp로 텍스트와 placeholder를 읽고,
#   Word(.docx) 회의록 양식으로 재구성한다.
# ============================================================

DEFAULT_MEETING_MINUTES_FIELDS = [
    "부서명",
    "건명",
    "장소",
    "소집 및 발안자",
    "의결 사항(요약)",
]

def extract_text_from_hwp_best_effort(hwp_path: str | Path) -> str:
    """
    HWP 바이너리에서 텍스트를 최대한 추출한다.
    hwp5txt CLI 또는 python -m hwp5.hwp5txt를 시도한다.
    """
    hwp_path = Path(hwp_path).resolve()

    candidates = [
        ["hwp5txt", str(hwp_path)],
        ["python", "-m", "hwp5.hwp5txt", str(hwp_path)],
    ]

    last_err = ""

    for cmd in candidates:
        try:
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=120,
            )

            if result.returncode == 0 and result.stdout.strip():
                return result.stdout

            last_err = (result.stderr or result.stdout or "")[-1500:]

        except Exception as e:
            last_err = str(e)

    raise RuntimeError(
        "HWP 텍스트 추출 실패. 이 HWP는 Ubuntu 서버에서 직접 읽지 못했습니다. "
        "가능하면 한글에서 HWPX로 저장한 뒤 업로드하세요. "
        f"detail={last_err}"
    )


def create_meeting_minutes_docx(output_path: str | Path, fields: dict) -> Path:
    """
    HWP 원본 서식 보존 대신, 회의록 표 양식을 DOCX로 재구성한다.
    """
    from pathlib import Path
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Malgun Gothic"
    normal_style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("회   의   록")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph("")

    dept = doc.add_paragraph()
    dept.add_run("부 서 명 : ").bold = True
    dept.add_run(str(fields.get("부서명", "미기재") or "미기재"))

    doc.add_paragraph("")

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    rows = [
        ("건     명", fields.get("건명", "")),
        ("진 행 시 간", fields.get("진행 시간", fields.get("진행시간", "    시      분부터\n    시      분까지"))),
        ("장     소", fields.get("장소", "")),
        ("소집 및 발안자", fields.get("소집 및 발안자", "")),
        ("의결 사항(요약)", fields.get("의결 사항(요약)", fields.get("의결사항", ""))),
    ]

    for i, (label, value) in enumerate(rows):
        left = table.cell(i, 0)
        right = table.cell(i, 1)

        left.text = label
        right.text = str(value or "미기재")

        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for para in left.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.bold = True

    doc.add_paragraph("")
    agree = doc.add_paragraph("위 결의사항에 모두 동의하며 연서 및 날인함.")
    agree.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")

    participants = fields.get("참석자", [])
    if isinstance(participants, str):
        participants = [
            x.strip()
            for x in participants.replace(",", "\n").splitlines()
            if x.strip()
        ]

    max_rows = max(10, len(participants) if isinstance(participants, list) else 0)

    for i in range(max_rows):
        name = participants[i] if isinstance(participants, list) and i < len(participants) else ""
        doc.add_paragraph(f" 참 석 자 : {name}                (인)")

    doc.save(str(output_path))
    return output_path


def generate_docx_from_hwp_text_template(
    template_path: str | Path,
    session_id: str,
    transcript: str,
    room_name: Optional[str],
    output_dir: str | Path,
    user_prompt: str = "",
) -> Dict[str, Any]:
    """
    .hwp 입력을 직접 HWPX로 변환하지 않고,
    텍스트/placeholder만 읽어 DOCX 회의록으로 재구성한다.
    """
    template_path = Path(template_path).resolve()
    output_dir = Path(output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    template_text = extract_text_from_hwp_best_effort(template_path)

    fields = extract_placeholder_fields(template_text)

    # HWP 텍스트 추출은 표 내부 placeholder 일부만 잡히는 경우가 많다.
    # 따라서 DOCX fallback에서는 기본 회의록 필드를 항상 보강한다.
    merged_fields = []
    for f in list(fields or []) + DEFAULT_MEETING_MINUTES_FIELDS:
        if f not in merged_fields:
            merged_fields.append(f)
    fields = merged_fields

    document_context = build_document_context_for_minutes(session_id, transcript)

    field_values = generate_field_values(
        fields,
        document_context,
        template_text=template_text,
        user_prompt=user_prompt,
    )

    # 기본 필드가 빠졌을 경우 보강
    # 중요: 원본 transcript가 아니라 정제된 document_context를 사용한다.
    for field in DEFAULT_MEETING_MINUTES_FIELDS:
        if field not in field_values:
            field_values[field] = build_field_fallback_value(field, document_context)
            field_values[field] = force_field_value_by_type(field, field_values[field], document_context)

    stem = template_path.stem
    output_docx = output_dir / f"{stem}_filled_{session_id[:8]}.docx"

    create_meeting_minutes_docx(output_docx, field_values)

    return {
        "ok": True,
        "templatePath": str(template_path),
        "roomName": room_name or "",
        "sessionId": session_id,
        "fields": fields,
        "fieldValues": field_values,
        "changedCount": len(field_values),
        "outputPath": str(output_docx),
        "outputDocxPath": str(output_docx),
        "outputHwpxPath": "",
        "model": HWP_REPORT_MODEL,
        "mode": "hwp_text_to_docx",
    }


# ============================================================
# HWPX 파싱/수정
# ============================================================

def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1].lower() if "}" in tag else tag.lower()


def register_xml_namespaces(xml_path: Path) -> None:
    try:
        for _, ns in __import__("xml.etree.ElementTree").ElementTree.iterparse(str(xml_path), events=("start-ns",)):
            prefix, uri = ns
            if prefix is None:
                prefix = ""
            __import__("xml.etree.ElementTree").ElementTree.register_namespace(prefix, uri)
    except Exception:
        pass


def iter_section_xml_files(extract_dir: Path) -> List[Path]:
    contents_dir = extract_dir / "Contents"
    if not contents_dir.exists():
        raise FileNotFoundError(f"HWPX Contents 폴더를 찾지 못했습니다: {contents_dir}")
    return sorted(contents_dir.glob("section*.xml"))


def get_text_nodes(elem: Any) -> List[Any]:
    return [node for node in elem.iter() if local_name(str(node.tag)) == "t"]


def paragraph_text(p_elem: Any) -> str:
    return "".join((node.text or "") for node in get_text_nodes(p_elem))


def load_paragraph_refs(section_files: List[Path]) -> Tuple[List[ParagraphRef], Dict[Path, Any]]:
    import xml.etree.ElementTree as ET

    refs: List[ParagraphRef] = []
    trees: Dict[Path, Any] = {}
    idx = 0
    for section_path in section_files:
        register_xml_namespaces(section_path)
        tree = ET.parse(str(section_path))
        trees[section_path] = tree
        root = tree.getroot()
        for elem in root.iter():
            if local_name(str(elem.tag)) == "p":
                text_nodes = get_text_nodes(elem)
                text = "".join((n.text or "") for n in text_nodes)
                refs.append(ParagraphRef(section_path, elem, text_nodes, text, idx))
                idx += 1
    return refs, trees


def extract_visible_template_text(paragraphs: List[ParagraphRef]) -> str:
    lines = []
    for ref in paragraphs:
        text = re.sub(r"\s+", " ", ref.text or "").strip()
        if text:
            lines.append(text)
    return "\n".join(lines)


def normalize_fields(fields: Iterable[str]) -> List[str]:
    seen = set()
    normalized: List[str] = []
    for field in fields:
        f = re.sub(r"[\s:：]+$", "", str(field or "").strip())
        f = re.sub(r"^[-•*\s]+", "", f)
        if not f or len(f) > 40:
            continue
        key = f.lower()
        if key not in seen:
            seen.add(key)
            normalized.append(f)
    return normalized


def extract_placeholder_fields(template_text: str) -> List[str]:
    fields: List[str] = []
    for pattern in PLACEHOLDER_PATTERNS:
        fields.extend(m.group(1).strip() for m in pattern.finditer(template_text or ""))
    return normalize_fields(fields)


def detect_fields_heuristically(template_text: str) -> List[str]:
    found = []
    for hint in COMMON_FIELD_HINTS:
        if re.search(re.escape(hint), template_text, flags=re.IGNORECASE):
            found.append(hint)

    # "항목명:" 패턴도 후보로 추출
    for line in template_text.splitlines():
        line = line.strip()
        m = re.match(r"^([가-힣A-Za-z0-9 ()/_-]{2,30})\s*[:：]", line)
        if m:
            found.append(m.group(1).strip())
    return normalize_fields(found)


def is_blank_like(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return True
    t2 = re.sub(r"[\s_\-—–=ㆍ·\.。:：/\\|]+", "", t)
    return len(t2) == 0


def replace_placeholders(paragraphs: List[ParagraphRef], field_values: Dict[str, str]) -> int:
    replaced = 0

    for ref in paragraphs:
        for node in ref.text_nodes:
            original = node.text or ""
            updated = original

            for field, value in field_values.items():
                value = str(value or "").strip()

                # 빈 값이면 placeholder를 지우지 않는다.
                if not value:
                    continue

                placeholder_candidates = [
                    "{{" + field + "}}",
                    "{{ " + field + " }}",
                    "[[" + field + "]]",
                    "[[ " + field + " ]]",
                    "${" + field + "}",
                    "${ " + field + " }",
                ]

                for ph in placeholder_candidates:
                    if ph in updated:
                        updated = updated.replace(ph, value)

            if updated != original:
                node.text = updated
                replaced += 1

    return replaced

def field_in_text(field: str, text: str) -> bool:
    compact_field = re.sub(r"\s+", "", field).lower()
    compact_text = re.sub(r"\s+", "", text or "").lower()
    return compact_field in compact_text


def replace_after_label_in_same_paragraph(ref: ParagraphRef, field: str, value: str) -> bool:
    """'회의 제목: ____' 같은 문단이면 콜론 뒤를 교체."""
    full = ref.text or ""
    if not field_in_text(field, full):
        return False
    pattern = re.compile(rf"({re.escape(field)}\s*[:：]\s*)(.*)$", flags=re.IGNORECASE)
    m = pattern.search(full)
    if not m:
        return False

    new_text = full[: m.start()] + m.group(1) + value
    if ref.text_nodes:
        ref.text_nodes[0].text = new_text
        for node in ref.text_nodes[1:]:
            node.text = ""
        return True
    return False


def find_next_blank_paragraph(paragraphs: List[ParagraphRef], start_index: int, max_scan: int = 8) -> Optional[ParagraphRef]:
    for ref in paragraphs[start_index + 1 : start_index + 1 + max_scan]:
        if ref.text_nodes and is_blank_like(ref.text):
            return ref
    return None


def append_value_to_label_paragraph(ref: ParagraphRef, field: str, value: str) -> bool:
    if not ref.text_nodes:
        return False
    # 마지막 텍스트 노드에 붙인다. 별도 blank cell이 없는 양식의 fallback.
    suffix = "\n" + value if value else ""
    ref.text_nodes[-1].text = (ref.text_nodes[-1].text or "") + suffix
    return True


def fill_by_label_proximity(paragraphs: List[ParagraphRef], field_values: Dict[str, str]) -> int:
    changed = 0
    used_blank_indexes = set()

    for field, value in field_values.items():
        if not value:
            continue
        candidates = [ref for ref in paragraphs if field_in_text(field, ref.text)]
        if not candidates:
            continue

        target_label = candidates[0]

        # 1) 같은 문단의 콜론 뒤 내용 교체
        if replace_after_label_in_same_paragraph(target_label, field, value):
            changed += 1
            continue

        # 2) 라벨 다음의 빈 문단/빈 셀에 삽입
        blank = find_next_blank_paragraph(paragraphs, target_label.index, max_scan=10)
        if blank is not None and blank.index not in used_blank_indexes:
            blank.text_nodes[0].text = value
            for node in blank.text_nodes[1:]:
                node.text = ""
            used_blank_indexes.add(blank.index)
            changed += 1
            continue

        # 3) fallback: 라벨 문단 자체에 줄바꿈 후 내용 추가
        if append_value_to_label_paragraph(target_label, field, value):
            changed += 1

    return changed


def write_xml_trees(trees: Dict[Path, Any]) -> None:
    for path, tree in trees.items():
        tree.write(str(path), encoding="utf-8", xml_declaration=True)


def unzip_hwpx(hwpx_path: Path, extract_dir: Path) -> None:
    with zipfile.ZipFile(str(hwpx_path), "r") as zipf:
        zipf.extractall(str(extract_dir))


def zip_hwpx_dir(extract_dir: Path, output_hwpx: Path) -> Path:
    output_hwpx.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(output_hwpx), "w", compression=zipfile.ZIP_DEFLATED) as zipf:
        for file_path in sorted(extract_dir.rglob("*")):
            if file_path.is_file():
                arcname = file_path.relative_to(extract_dir).as_posix()
                zipf.write(str(file_path), arcname)
    return output_hwpx



# ============================================================
# FORCE DOCX ONLY MODE
# - 입력: .hwp / .hwpx / .docx
# - 출력: 무조건 .docx
# - HWP/HWPX output 변환은 사용하지 않는다.
# ============================================================

FORCE_DOCX_ONLY = True

def force_extract_text_from_hwp(hwp_path: str | Path) -> str:
    hwp_path = Path(hwp_path).resolve()

    candidates = [
        ["hwp5txt", str(hwp_path)],
        ["python", "-m", "hwp5.hwp5txt", str(hwp_path)],
    ]

    last_err = ""

    for cmd in candidates:
        try:
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=120,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
            last_err = (result.stderr or result.stdout or "")[-1500:]
        except Exception as e:
            last_err = str(e)

    raise RuntimeError(
        "HWP 텍스트 추출 실패. 이 HWP는 서버에서 직접 읽지 못했습니다. "
        "가능하면 HWPX 또는 DOCX 양식으로 저장해서 업로드하세요. "
        f"detail={last_err}"
    )


def force_extract_text_from_docx(docx_path: str | Path) -> str:
    from docx import Document

    doc = Document(str(docx_path))
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def force_replace_docx_paragraph(paragraph, field_values: dict) -> int:
    original = paragraph.text or ""
    updated = original

    for field, value in field_values.items():
        value = str(value or "").strip()
        if not value:
            continue

        placeholders = [
            "{{" + field + "}}",
            "{{ " + field + " }}",
            "[[" + field + "]]",
            "[[ " + field + " ]]",
            "${" + field + "}",
            "${ " + field + " }",
        ]

        for ph in placeholders:
            updated = updated.replace(ph, value)

    if updated == original:
        return 0

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)

    return 1


def force_fill_docx_template(template_path: str | Path, output_path: str | Path, field_values: dict) -> int:
    from docx import Document

    template_path = Path(template_path).resolve()
    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))
    changed = 0

    for p in doc.paragraphs:
        changed += force_replace_docx_paragraph(p, field_values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    changed += force_replace_docx_paragraph(p, field_values)

    doc.save(str(output_path))
    return changed


def force_create_meeting_minutes_docx(output_path: str | Path, fields: dict) -> Path:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("회   의   록")
    r.bold = True
    r.font.size = Pt(20)

    doc.add_paragraph("")

    dept = doc.add_paragraph()
    dept.add_run("부 서 명 : ").bold = True
    dept.add_run(str(fields.get("부서명", "미기재") or "미기재"))

    doc.add_paragraph("")

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    rows = [
        ("건     명", fields.get("건명", "회의 내용 기반 요약")),
        ("진 행 시 간", fields.get("진행 시간", fields.get("진행시간", "    시      분부터\n    시      분까지"))),
        ("장     소", fields.get("장소", "미기재")),
        ("소집 및 발안자", fields.get("소집 및 발안자", "미기재")),
        ("의결 사항(요약)", fields.get("의결 사항(요약)", fields.get("의결사항", "미기재"))),
    ]

    for i, (label, value) in enumerate(rows):
        left = table.cell(i, 0)
        right = table.cell(i, 1)
        left.text = label
        right.text = str(value or "미기재")

        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for para in left.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("위 결의사항에 모두 동의하며 연서 및 날인함.")
    doc.add_paragraph("")

    participants = fields.get("참석자", [])
    if isinstance(participants, str):
        participants = [
            x.strip()
            for x in participants.replace(",", "\n").splitlines()
            if x.strip()
        ]

    if not isinstance(participants, list):
        participants = []

    for i in range(max(10, len(participants))):
        name = participants[i] if i < len(participants) else ""
        doc.add_paragraph(f" 참 석 자 : {name}                (인)")

    doc.save(str(output_path))
    return output_path


def force_get_template_text(template_path: str | Path) -> str:
    template_path = Path(template_path).resolve()
    ext = template_path.suffix.lower()

    if ext == ".docx":
        return force_extract_text_from_docx(template_path)

    if ext == ".hwp":
        return force_extract_text_from_hwp(template_path)

    if ext == ".hwpx":
        with tempfile.TemporaryDirectory(prefix="force_hwpx_read_") as tmp:
            work_dir = Path(tmp)
            extract_dir = work_dir / "unzipped"
            unzip_hwpx(template_path, extract_dir)
            section_files = iter_section_xml_files(extract_dir)
            paragraphs, _trees = load_paragraph_refs(section_files)
            return extract_visible_template_text(paragraphs)

    raise ValueError(f"지원하지 않는 템플릿 형식입니다: {ext}")


def force_generate_docx_only_report(
    template_path: str | Path,
    session_id: str,
    room_name: Optional[str],
    output_dir: str | Path,
    use_llm_field_detection: bool = True,
    user_prompt: str = "",
) -> Dict[str, Any]:
    template_path = Path(template_path).resolve()
    output_dir = Path(output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    transcript = read_session_transcript(room_name, session_id)
    if not transcript:
        raise ValueError("해당 세션의 회의 STT가 비어 있습니다.")

    document_context = build_document_context_for_minutes(session_id, transcript)

    template_text = force_get_template_text(template_path)

    fields = extract_placeholder_fields(template_text)

    if not fields and use_llm_field_detection:
        try:
            fields = detect_fields_with_qwen(template_text)
        except Exception as e:
            print(f"[WARN] force docx field detection failed: {e}")

    default_fields = globals().get("DEFAULT_MEETING_MINUTES_FIELDS", [
        "부서명",
        "건명",
        "장소",
        "소집 및 발안자",
        "의결 사항(요약)",
    ])

    merged_fields = []
    for f in list(fields or []) + list(default_fields):
        if f not in merged_fields:
            merged_fields.append(f)

    fields = merged_fields

    field_values = generate_field_values(
        fields,
        document_context,
        template_text=template_text,
        user_prompt=user_prompt,
    )

    for field in default_fields:
        if field not in field_values:
            field_values[field] = build_field_fallback_value(field, document_context)
            field_values[field] = force_field_value_by_type(field, field_values[field], document_context)

    stem = template_path.stem
    output_docx = output_dir / f"{stem}_filled_{session_id[:8]}.docx"

    changed = 0

    if template_path.suffix.lower() == ".docx":
        changed = force_fill_docx_template(template_path, output_docx, field_values)

    if template_path.suffix.lower() != ".docx" or changed == 0:
        force_create_meeting_minutes_docx(output_docx, field_values)
        changed = len(field_values)

    return {
        "ok": True,
        "templatePath": str(template_path),
        "roomName": room_name or "",
        "sessionId": session_id,
        "fields": fields,
        "fieldValues": field_values,
        "changedCount": changed,
        "outputPath": str(output_docx),
        "outputDocxPath": str(output_docx),
        "outputHwpxPath": "",
        "model": HWP_REPORT_MODEL,
        "mode": "force_docx_only",
    }


# ============================================================
# 전체 파이프라인
# ============================================================

def prepare_template_as_hwpx(template_path: Path, work_dir: Path) -> Path:
    ext = template_path.suffix.lower()
    temp_hwpx = work_dir / f"template_{uuid.uuid4().hex}.hwpx"

    if ext == ".hwpx":
        shutil.copy2(str(template_path), str(temp_hwpx))
        return temp_hwpx
    if ext == ".hwp":
        return convert_hwp_to_hwpx(template_path, temp_hwpx)
    raise ValueError(f"지원하지 않는 양식 파일 형식입니다: {ext}")


def determine_fields(template_text: str, use_llm_field_detection: bool = True) -> List[str]:
    placeholder_fields = extract_placeholder_fields(template_text)
    if placeholder_fields:
        return placeholder_fields

    if use_llm_field_detection:
        try:
            fields = detect_fields_with_qwen(template_text)
            if fields:
                return fields
        except Exception as e:
            print(f"[WARN] QWEN 필드 감지 실패, 휴리스틱으로 전환: {e}")

    return detect_fields_heuristically(template_text)


def generate_report_from_hwp_template(
    template_path: str | Path,
    session_id: str,
    room_name: Optional[str] = None,
    output_format: str = "hwp",
    output_dir: str | Path | None = None,
    use_llm_field_detection: bool = True,
    user_prompt: str = "",
) -> Dict[str, Any]:
    """
    HWP/HWPX 양식 + 회의 세션 STT → 채워진 회의요약본 생성.

    output_format: "hwp", "hwpx", 또는 "docx"
    """
    template_path = Path(template_path).resolve()
    if not template_path.exists():
        raise FileNotFoundError(f"양식 파일이 없습니다: {template_path}")

    # 현재 서비스는 안정성을 위해 무조건 DOCX만 출력한다.
    output_format = "docx"

    transcript = read_session_transcript(room_name, session_id)
    if not transcript:
        raise ValueError("해당 세션의 회의 STT가 비어 있습니다.")

    out_dir = Path(output_dir).resolve() if output_dir else OUTPUT_ROOT / datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir.mkdir(parents=True, exist_ok=True)

    # FORCE DOCX ONLY: hwp/hwpx/docx 입력 모두 docx로만 생성한다.
    return force_generate_docx_only_report(
        template_path=template_path,
        session_id=session_id,
        room_name=room_name,
        output_dir=out_dir,
        use_llm_field_detection=use_llm_field_detection,
        user_prompt=user_prompt,
    )

    # Linux/Ubuntu fallback:
    # .hwp 입력 + .docx 출력은 HWPX 변환 없이 텍스트 추출 후 DOCX로 재구성한다.
    if template_path.suffix.lower() == ".hwp" and output_format == "docx":
        return generate_docx_from_hwp_text_template(
            template_path=template_path,
            session_id=session_id,
            transcript=transcript,
            room_name=room_name,
            output_dir=out_dir,
            user_prompt=user_prompt,
        )

    with tempfile.TemporaryDirectory(prefix="hwp_template_fill_") as tmp:
        work_dir = Path(tmp)
        template_hwpx = prepare_template_as_hwpx(template_path, work_dir)
        extract_dir = work_dir / "unzipped"
        unzip_hwpx(template_hwpx, extract_dir)

        section_files = iter_section_xml_files(extract_dir)
        paragraphs, trees = load_paragraph_refs(section_files)
        template_text = extract_visible_template_text(paragraphs)

        fields = determine_fields(template_text, use_llm_field_detection=use_llm_field_detection)
        if not fields:
            raise ValueError(
                "양식에서 채울 항목을 찾지 못했습니다. "
                "양식에 {{회의 제목}}, {{회의 내용}} 같은 placeholder를 넣으면 가장 안정적입니다."
            )

        document_context = build_document_context_for_minutes(session_id, transcript)

        field_values = generate_field_values(
            fields,
            document_context,
            template_text=template_text,
            user_prompt=user_prompt,
        )

        changed = replace_placeholders(paragraphs, field_values)

        # 표 양식에서는 label proximity 방식이 잘못된 칸에 값을 넣을 수 있으므로 기본적으로 막는다.
        # 꼭 기존 방식도 쓰고 싶으면 환경변수 HWP_ALLOW_LABEL_PROXIMITY=1 일 때만 허용한다.
        if changed == 0 and os.getenv("HWP_ALLOW_LABEL_PROXIMITY", "0") == "1":
            changed = fill_by_label_proximity(paragraphs, field_values)

        if changed == 0:
            raise RuntimeError(
                "양식에 내용을 넣을 정확한 위치를 찾지 못했습니다. "
                "HWP 양식의 입력 위치에 {{부서명}}, {{건명}}, {{장소}}, "
                "{{소집 및 발안자}}, {{의결 사항(요약)}} 같은 placeholder를 넣어 주세요."
            )
        write_xml_trees(trees)

        stem = template_path.stem
        output_hwpx = out_dir / f"{stem}_filled_{session_id[:8]}.hwpx"
        zip_hwpx_dir(extract_dir, output_hwpx)

        if output_format == "hwpx":
            final_path = output_hwpx
        elif output_format == "hwp":
            final_path = out_dir / f"{stem}_filled_{session_id[:8]}.hwp"
            convert_hwpx_to_hwp(output_hwpx, final_path)
        elif output_format == "docx":
            # HWPX 입력을 DOCX로 요청한 경우에는 원본 HWPX 서식 변환 대신
            # 추출된 필드값으로 Word 회의록 양식을 재구성한다.
            final_path = out_dir / f"{stem}_filled_{session_id[:8]}.docx"
            create_meeting_minutes_docx(final_path, field_values)
        else:
            raise ValueError(f"지원하지 않는 output_format입니다: {output_format}")

    return {
        "ok": True,
        "templatePath": str(template_path),
        "roomName": room_name or "",
        "sessionId": session_id,
        "fields": fields,
        "fieldValues": field_values,
        "changedCount": changed,
        "outputPath": str(final_path),
        "outputHwpxPath": str(output_hwpx),
        "model": HWP_REPORT_MODEL,
    }


# ============================================================
# FastAPI Router
# ============================================================

if APIRouter is not None:
    router = APIRouter(prefix="/meeting-report-template", tags=["Meeting Report Template"])

    @router.post("/generate")
    async def generate_meeting_report_template(
        template: UploadFile = File(...),
        sessionId: str = Form(...),
        roomName: str = Form(""),
        outputFormat: str = Form("docx"),
        useLlmFieldDetection: bool = Form(True),
        userPrompt: str = Form(""),
    ):
        """
        multipart/form-data:
        - template: .hwp 또는 .hwpx 파일
        - sessionId: 회의 세션 ID
        - roomName: 룸 이름. 비워도 sessionId로 backend/data 내부 검색 시도
        - outputFormat: hwp, hwpx 또는 docx
        """
        ext = Path(template.filename or "").suffix.lower()
        if ext not in {".hwp", ".hwpx", ".docx"}:
            raise HTTPException(status_code=400, detail="template은 .hwp, .hwpx 또는 .docx만 가능합니다.")

        req_id = uuid.uuid4().hex
        upload_dir = OUTPUT_ROOT / "uploads" / req_id
        upload_dir.mkdir(parents=True, exist_ok=True)
        saved_template = upload_dir / f"template{ext}"

        try:
            content = await template.read()
            saved_template.write_bytes(content)

            result = generate_report_from_hwp_template(
                template_path=saved_template,
                session_id=sessionId,
                room_name=roomName or None,
                output_format=outputFormat,
                output_dir=OUTPUT_ROOT / req_id,
                use_llm_field_detection=useLlmFieldDetection,
                user_prompt=userPrompt,
            )
            output_path = Path(result["outputPath"])
            if not output_path.exists():
                raise RuntimeError(f"생성 결과 파일을 찾을 수 없습니다: {output_path}")

            media_type = "application/octet-stream"
            if output_path.suffix.lower() == ".docx":
                media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif output_path.suffix.lower() == ".hwpx":
                media_type = "application/zip"

            return FileResponse(
                path=str(output_path),
                filename=output_path.name,
                media_type=media_type,
            )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"회의요약본 양식 작성 실패: {str(e)}")
else:
    router = None


# ============================================================
# CLI
# ============================================================

def main() -> None:
    parser = argparse.ArgumentParser(description="HWP 회의요약본 양식을 회의 세션 STT 기반으로 자동 작성")
    parser.add_argument("--template", required=True, help="회의요약본 양식 .hwp 또는 .hwpx 경로")
    parser.add_argument("--session", required=True, help="session_id")
    parser.add_argument("--room", default="", help="room_name. 비우면 session_id로 backend/data 전체 검색")
    parser.add_argument("--format", default="docx", choices=["docx"], help="출력 형식")
    parser.add_argument("--output-dir", default="", help="출력 폴더")
    parser.add_argument("--no-llm-field-detect", action="store_true", help="양식 필드 감지에 GEMMA를 쓰지 않고 휴리스틱만 사용")
    args = parser.parse_args()

    kill_hwp_processes()
    result = generate_report_from_hwp_template(
        template_path=args.template,
        session_id=args.session,
        room_name=args.room or None,
        output_format=args.format,
        output_dir=args.output_dir or None,
        use_llm_field_detection=not args.no_llm_field_detect,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
