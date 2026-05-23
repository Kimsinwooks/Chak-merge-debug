import argparse
import sqlite3
from pathlib import Path

from docx import Document

from hwp_template_report_api import (
    DATA_DIR,
    read_session_transcript,
    read_cached_meeting_report_text,
    build_document_context_for_minutes,
    generate_report_from_hwp_template,
)


BAD_RAW_PHRASES = [
    "안되고 목요일에 하든가",
    "좋아요 부분이 있어요",
    "제가 부탁이니까",
    "뭐 할건지 모르고",
    "답변이 안되잖아",
    "네 네",
    "익명1: 네",
]


def find_latest_session_id():
    db = DATA_DIR / "meeting_app.sqlite3"
    if not db.exists():
        raise RuntimeError(f"DB not found: {db}")

    conn = sqlite3.connect(db)
    cur = conn.cursor()

    candidates = []

    for table in ["meeting_report_cache", "transcript_lines", "meeting_sessions"]:
        row = cur.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
            (table,),
        ).fetchone()
        if not row:
            continue

        cols = [c[1] for c in cur.execute(f"PRAGMA table_info({table})")]
        if "session_id" not in cols:
            continue

        try:
            rows = cur.execute(
                f"""
                SELECT session_id, COUNT(*) AS n
                FROM {table}
                WHERE session_id IS NOT NULL AND session_id != ''
                GROUP BY session_id
                ORDER BY rowid DESC
                LIMIT 10
                """
            ).fetchall()
            for sid, n in rows:
                candidates.append((table, sid, n))
        except Exception:
            pass

    conn.close()

    if not candidates:
        raise RuntimeError("No session_id found in DB")

    print("[CANDIDATE SESSIONS]")
    for table, sid, n in candidates[:10]:
        print(f"  table={table:22s} session_id={sid} count={n}")

    return candidates[0][1]


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True, help=".hwp 또는 .hwpx 양식 경로")
    parser.add_argument("--session-id", default="", help="비우면 DB에서 최신 session_id 자동 선택")
    parser.add_argument("--room-name", default="", help="룸 이름. 비워도 됨")
    parser.add_argument("--output-format", default="docx", choices=["docx", "hwpx", "hwp"])
    parser.add_argument("--user-prompt", default="공식 회의록 문체로 작성. STT 원문 말투를 복사하지 말 것.")
    args = parser.parse_args()

    template = Path(args.template).expanduser().resolve()
    if not template.exists():
        raise RuntimeError(f"template not found: {template}")

    session_id = args.session_id.strip() or find_latest_session_id()
    room_name = args.room_name.strip() or None

    print("\n[TEST CONFIG]")
    print("template     =", template)
    print("session_id   =", session_id)
    print("room_name    =", room_name)
    print("output_format=", args.output_format)

    print("\n[1] transcript check")
    transcript = read_session_transcript(room_name, session_id)
    print("transcript chars =", len(transcript))
    print("transcript preview:")
    print(transcript[:800])

    print("\n[2] cached report check")
    cached = read_cached_meeting_report_text(session_id)
    print("cached report chars =", len(cached))
    if cached:
        print("cached report preview:")
        print(cached[:1200])
    else:
        print("[WARN] cached report not found. 문서 생성이 transcript fallback으로 갈 수 있음.")

    print("\n[3] document context selection")
    context = build_document_context_for_minutes(session_id, transcript)
    print("context chars =", len(context))
    if cached and context.strip() == cached.strip():
        print("[OK] using cached meeting report context")
    elif cached and cached[:100] in context:
        print("[OK] context includes cached report")
    else:
        print("[WARN] cached report is not used. fallback transcript may be used.")

    print("\n[4] generate document")
    out_dir = DATA_DIR / "debug_hwp_docx_test"
    result = generate_report_from_hwp_template(
        template_path=template,
        session_id=session_id,
        room_name=room_name,
        output_format=args.output_format,
        output_dir=out_dir,
        use_llm_field_detection=False,
        user_prompt=args.user_prompt,
    )

    print("result:")
    for k, v in result.items():
        if k == "fieldValues":
            print("fieldValues:")
            for fk, fv in v.items():
                print(f"  - {fk}: {str(fv)[:300]}")
        else:
            print(f"{k}: {v}")

    output_path = Path(result["outputPath"])
    print("\n[5] output file check")
    print("output_path =", output_path)
    print("exists      =", output_path.exists())
    print("size        =", output_path.stat().st_size if output_path.exists() else 0)

    if args.output_format == "docx":
        print("\n[6] docx text inspection")
        docx_text = read_docx_text(output_path)
        print("docx chars =", len(docx_text))
        print("docx preview:")
        print(docx_text[:2000])

        print("\n[7] quality checks")

        required_keywords = ["회", "의", "록", "건", "명", "장", "소", "의결"]
        missing = [x for x in required_keywords if x not in docx_text]
        if missing:
            print("[WARN] missing expected keywords:", missing)
        else:
            print("[OK] basic meeting-minutes keywords found")

        bad_found = [p for p in BAD_RAW_PHRASES if p in docx_text]
        if bad_found:
            print("[FAIL] raw STT-like phrases found:")
            for p in bad_found:
                print("  -", p)
        else:
            print("[OK] no known bad raw STT phrases found")

        if "## 회의 개요" in docx_text or "주요 내용" in docx_text:
            print("[WARN] markdown heading leaked into DOCX. 필드 생성 prompt 조정 필요.")
        else:
            print("[OK] markdown heading not directly leaked")

    print("\n[DONE]")


if __name__ == "__main__":
    main()
